# retrieval.py

import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS

# Load the DOCX file
def load_docx(file_path):
    doc = docx.Document(file_path)
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"
    return full_text

file_path = "FAQ_Flipkart.docx"  
doc_text = load_docx(file_path)

# Split text into chunks
text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
chunks = text_splitter.split_text(doc_text)
documents = [Document(page_content=chunk) for chunk in chunks]

# Initialize vectorstore
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
vectorstore = FAISS.from_documents(documents, embeddings)

# Retrieval function
def retrieve_top_chunks(query: str, k: int = 3):
    results = vectorstore.similarity_search(query, k=k)
    return [doc.page_content for doc in results]
